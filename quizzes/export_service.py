"""
Advanced quiz export functionality supporting multiple formats
"""
import os
import json
import csv
import io
from datetime import datetime
from typing import Dict, List, Any, Optional
import logging

from django.conf import settings
from django.http import HttpResponse, Http404
from django.template.loader import render_to_string
from django.utils import timezone
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404

logger = logging.getLogger(__name__)

class QuizExportService:
    """Comprehensive quiz export service"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'json', 'csv', 'xlsx', 'html']
    
    def export_quiz(self, quiz, format_type: str, user, **options) -> Dict[str, Any]:
        """Export quiz in specified format"""
        if format_type not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format_type}")
        
        # Prepare quiz data
        quiz_data = self._prepare_quiz_data(quiz, **options)
        
        # Generate export based on format
        export_method = getattr(self, f'_export_to_{format_type}')
        result = export_method(quiz_data, quiz, **options)
        
        # Track export
        self._track_export(quiz, user, format_type, **options)
        
        return result
    
    def _prepare_quiz_data(self, quiz, **options) -> Dict[str, Any]:
        """Prepare comprehensive quiz data for export"""
        include_answers = options.get('include_answers', True)
        include_explanations = options.get('include_explanations', True)
        include_analytics = options.get('include_analytics', False)
        include_metadata = options.get('include_metadata', True)
        
        # Basic quiz information
        quiz_data = {
            'title': quiz.title,
            'description': quiz.description,
            'category': quiz.category.name if quiz.category else 'Uncategorized',
            'difficulty': quiz.get_difficulty_display(),
            'created_by': quiz.created_by.get_full_name() or quiz.created_by.username,
            'created_at': quiz.created_at.isoformat(),
            'total_questions': quiz.questions.count(),
            'estimated_time': f"{quiz.time_limit} minutes" if quiz.time_limit else "No time limit",
            'questions': []
        }
        
        # Add metadata if requested
        if include_metadata:
            quiz_data['metadata'] = {
                'quiz_type': quiz.get_quiz_type_display(),
                'max_attempts': quiz.max_attempts,
                'points_reward': quiz.points_reward,
                'tags': quiz.tags or [],
                'is_published': quiz.is_published,
                'is_featured': quiz.is_featured,
            }
        
        # Add analytics if requested
        if include_analytics and hasattr(quiz, 'analytics'):
            analytics = quiz.analytics
            quiz_data['analytics'] = {
                'total_attempts': analytics.total_attempts,
                'unique_users': analytics.unique_users,
                'average_score': analytics.average_score,
                'completion_rate': analytics.completion_rate,
                'average_time_taken': analytics.average_time_taken,
                'difficulty_rating': analytics.difficulty_rating,
            }
        
        # Process questions
        for question in quiz.questions.all().order_by('order'):
            question_data = {
                'id': question.id,
                'text': question.text,
                'type': question.get_question_type_display(),
                'order': question.order,
                'points': question.points,
                'difficulty': question.get_difficulty_display(),
            }
            
            # Add options for multiple choice questions
            if question.question_type == 'MULTIPLE_CHOICE':
                question_data['options'] = {
                    'A': question.option_a,
                    'B': question.option_b,
                    'C': question.option_c,
                    'D': question.option_d,
                }
                
                if include_answers:
                    question_data['correct_option'] = question.correct_option
            
            elif question.question_type == 'TRUE_FALSE':
                question_data['options'] = ['True', 'False']
                if include_answers:
                    question_data['correct_answer'] = 'True' if question.correct_option == 'A' else 'False'
            
            elif question.question_type == 'SHORT_ANSWER':
                if include_answers:
                    question_data['correct_answer'] = question.correct_answer
            
            # Add explanation if requested
            if include_explanations and question.explanation:
                question_data['explanation'] = question.explanation
            
            # Add enhanced data if available
            if hasattr(question, 'enhanced'):
                enhanced = question.enhanced
                question_data['enhanced'] = {
                    'cognitive_level': enhanced.get_cognitive_level_display(),
                    'estimated_time': f"{enhanced.estimated_time} seconds",
                    'topic_tags': enhanced.topic_tags,
                }
                
                if include_analytics:
                    question_data['analytics'] = {
                        'correct_rate': enhanced.correct_rate,
                        'average_time_taken': enhanced.average_time_taken,
                        'skip_rate': enhanced.skip_rate,
                    }
            
            quiz_data['questions'].append(question_data)
        
        return quiz_data
    
    def _export_to_pdf(self, quiz_data: Dict, quiz, **options) -> Dict[str, Any]:
        """Export quiz to PDF format"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            
            # Create buffer
            buffer = io.BytesIO()
            
            # Create PDF document
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1  # Center
            )
            
            # Build story
            story = []
            
            # Title
            story.append(Paragraph(quiz_data['title'], title_style))
            story.append(Spacer(1, 12))
            
            # Quiz info
            info_data = [
                ['Category:', quiz_data['category']],
                ['Difficulty:', quiz_data['difficulty']],
                ['Total Questions:', str(quiz_data['total_questions'])],
                ['Created by:', quiz_data['created_by']],
                ['Created on:', quiz_data['created_at'][:10]],
            ]
            
            if quiz_data.get('estimated_time'):
                info_data.append(['Estimated Time:', quiz_data['estimated_time']])
            
            info_table = Table(info_data, colWidths=[2*inch, 4*inch])
            info_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ]))
            
            story.append(info_table)
            story.append(Spacer(1, 20))
            
            # Description
            if quiz_data.get('description'):
                story.append(Paragraph('<b>Description:</b>', styles['Heading3']))
                story.append(Paragraph(quiz_data['description'], styles['Normal']))
                story.append(Spacer(1, 15))
            
            # Questions
            story.append(Paragraph('Questions', styles['Heading2']))
            story.append(Spacer(1, 12))
            
            for i, question in enumerate(quiz_data['questions'], 1):
                # Question number and text
                q_text = f"<b>Question {i}:</b> {question['text']}"
                story.append(Paragraph(q_text, styles['Normal']))
                story.append(Spacer(1, 6))
                
                # Options for multiple choice
                if question.get('options') and isinstance(question['options'], dict):
                    for option_key, option_text in question['options'].items():
                        option_style = styles['Normal']
                        if options.get('include_answers') and option_key == question.get('correct_option'):
                            option_style = ParagraphStyle(
                                'CorrectOption',
                                parent=styles['Normal'],
                                textColor=colors.green,
                                fontName='Helvetica-Bold'
                            )
                        
                        story.append(Paragraph(f"{option_key}. {option_text}", option_style))
                        story.append(Spacer(1, 3))
                
                # Correct answer for other question types
                elif options.get('include_answers') and question.get('correct_answer'):
                    story.append(Paragraph(f"<b>Answer:</b> {question['correct_answer']}", styles['Normal']))
                    story.append(Spacer(1, 6))
                
                # Explanation
                if options.get('include_explanations') and question.get('explanation'):
                    explanation_style = ParagraphStyle(
                        'Explanation',
                        parent=styles['Normal'],
                        textColor=colors.blue,
                        leftIndent=20
                    )
                    story.append(Paragraph(f"<b>Explanation:</b> {question['explanation']}", explanation_style))
                    story.append(Spacer(1, 6))
                
                story.append(Spacer(1, 15))
            
            # Build PDF
            doc.build(story)
            
            # Get PDF data
            pdf_data = buffer.getvalue()
            buffer.close()
            
            return {
                'content': pdf_data,
                'content_type': 'application/pdf',
                'filename': f"{quiz.slug}_quiz.pdf"
            }
            
        except ImportError:
            logger.error("ReportLab not installed for PDF export")
            raise ValueError("PDF export requires ReportLab package")
        except Exception as e:
            logger.error(f"PDF export failed: {str(e)}")
            raise
    
    def _export_to_docx(self, quiz_data: Dict, quiz, **options) -> Dict[str, Any]:
        """Export quiz to Word document format"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading(quiz_data['title'], 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add quiz information
            doc.add_heading('Quiz Information', level=1)
            
            info_table = doc.add_table(rows=5, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('Category', quiz_data['category']),
                ('Difficulty', quiz_data['difficulty']),
                ('Total Questions', str(quiz_data['total_questions'])),
                ('Created by', quiz_data['created_by']),
                ('Created on', quiz_data['created_at'][:10]),
            ]
            
            for i, (label, value) in enumerate(info_data):
                info_table.cell(i, 0).text = label
                info_table.cell(i, 1).text = value
            
            # Add description
            if quiz_data.get('description'):
                doc.add_heading('Description', level=1)
                doc.add_paragraph(quiz_data['description'])
            
            # Add questions
            doc.add_heading('Questions', level=1)
            
            for i, question in enumerate(quiz_data['questions'], 1):
                # Question heading
                doc.add_heading(f'Question {i}', level=2)
                doc.add_paragraph(question['text'])
                
                # Options
                if question.get('options') and isinstance(question['options'], dict):
                    for option_key, option_text in question['options'].items():
                        p = doc.add_paragraph()
                        if options.get('include_answers') and option_key == question.get('correct_option'):
                            run = p.add_run(f"{option_key}. {option_text}")
                            run.bold = True
                        else:
                            p.add_run(f"{option_key}. {option_text}")
                
                # Correct answer for other types
                elif options.get('include_answers') and question.get('correct_answer'):
                    p = doc.add_paragraph()
                    p.add_run('Answer: ').bold = True
                    p.add_run(question['correct_answer'])
                
                # Explanation
                if options.get('include_explanations') and question.get('explanation'):
                    p = doc.add_paragraph()
                    p.add_run('Explanation: ').bold = True
                    p.add_run(question['explanation'])
                
                doc.add_paragraph()  # Add space
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            
            return {
                'content': buffer.getvalue(),
                'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'filename': f"{quiz.slug}_quiz.docx"
            }
            
        except ImportError:
            logger.error("python-docx not installed for DOCX export")
            raise ValueError("DOCX export requires python-docx package")
        except Exception as e:
            logger.error(f"DOCX export failed: {str(e)}")
            raise
    
    def _export_to_json(self, quiz_data: Dict, quiz, **options) -> Dict[str, Any]:
        """Export quiz to JSON format"""
        json_content = json.dumps(quiz_data, indent=2, ensure_ascii=False)
        
        return {
            'content': json_content.encode('utf-8'),
            'content_type': 'application/json',
            'filename': f"{quiz.slug}_quiz.json"
        }
    
    def _export_to_csv(self, quiz_data: Dict, quiz, **options) -> Dict[str, Any]:
        """Export quiz to CSV format"""
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Header
        headers = ['Question Number', 'Question Text', 'Question Type', 'Points']
        
        if quiz_data['questions'] and quiz_data['questions'][0].get('options'):
            headers.extend(['Option A', 'Option B', 'Option C', 'Option D'])
            if options.get('include_answers'):
                headers.append('Correct Option')
        
        if options.get('include_answers'):
            headers.append('Correct Answer')
        
        if options.get('include_explanations'):
            headers.append('Explanation')
        
        writer.writerow(headers)
        
        # Data rows
        for i, question in enumerate(quiz_data['questions'], 1):
            row = [
                i,
                question['text'],
                question['type'],
                question.get('points', 10)
            ]
            
            # Add options if they exist
            if question.get('options') and isinstance(question['options'], dict):
                options_dict = question['options']
                row.extend([
                    options_dict.get('A', ''),
                    options_dict.get('B', ''),
                    options_dict.get('C', ''),
                    options_dict.get('D', '')
                ])
                
                if options.get('include_answers'):
                    row.append(question.get('correct_option', ''))
            
            # Add correct answer
            if options.get('include_answers') and question.get('correct_answer'):
                row.append(question['correct_answer'])
            elif options.get('include_answers'):
                row.append('')
            
            # Add explanation
            if options.get('include_explanations'):
                row.append(question.get('explanation', ''))
            
            writer.writerow(row)
        
        return {
            'content': output.getvalue().encode('utf-8'),
            'content_type': 'text/csv',
            'filename': f"{quiz.slug}_quiz.csv"
        }
    
    def _export_to_xlsx(self, quiz_data: Dict, quiz, **options) -> Dict[str, Any]:
        """Export quiz to Excel format"""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment
            
            wb = openpyxl.Workbook()
            
            # Quiz Info Sheet
            ws_info = wb.active
            ws_info.title = "Quiz Info"
            
            # Style for headers
            header_font = Font(bold=True, size=12)
            header_fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
            
            # Add quiz information
            info_data = [
                ['Quiz Information', ''],
                ['Title', quiz_data['title']],
                ['Category', quiz_data['category']],
                ['Difficulty', quiz_data['difficulty']],
                ['Total Questions', quiz_data['total_questions']],
                ['Created by', quiz_data['created_by']],
                ['Created on', quiz_data['created_at'][:10]],
            ]
            
            if quiz_data.get('description'):
                info_data.append(['Description', quiz_data['description']])
            
            for row_num, (label, value) in enumerate(info_data, 1):
                ws_info.cell(row=row_num, column=1, value=label).font = header_font
                ws_info.cell(row=row_num, column=2, value=value)
            
            # Questions Sheet
            ws_questions = wb.create_sheet(title="Questions")
            
            # Headers
            headers = ['#', 'Question', 'Type', 'Points']
            
            if quiz_data['questions'] and quiz_data['questions'][0].get('options'):
                headers.extend(['Option A', 'Option B', 'Option C', 'Option D'])
                if options.get('include_answers'):
                    headers.append('Correct')
            
            if options.get('include_answers'):
                headers.append('Answer')
            
            if options.get('include_explanations'):
                headers.append('Explanation')
            
            # Write headers
            for col_num, header in enumerate(headers, 1):
                cell = ws_questions.cell(row=1, column=col_num, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center')
            
            # Write questions
            for row_num, question in enumerate(quiz_data['questions'], 2):
                ws_questions.cell(row=row_num, column=1, value=row_num - 1)
                ws_questions.cell(row=row_num, column=2, value=question['text'])
                ws_questions.cell(row=row_num, column=3, value=question['type'])
                ws_questions.cell(row=row_num, column=4, value=question.get('points', 10))
                
                col_offset = 5
                
                # Add options if they exist
                if question.get('options') and isinstance(question['options'], dict):
                    options_dict = question['options']
                    for i, option_key in enumerate(['A', 'B', 'C', 'D']):
                        ws_questions.cell(
                            row=row_num, 
                            column=col_offset + i, 
                            value=options_dict.get(option_key, '')
                        )
                    
                    col_offset += 4
                    
                    if options.get('include_answers'):
                        correct_cell = ws_questions.cell(
                            row=row_num, 
                            column=col_offset, 
                            value=question.get('correct_option', '')
                        )
                        correct_cell.font = Font(bold=True, color="008000")
                        col_offset += 1
                
                # Add correct answer
                if options.get('include_answers') and question.get('correct_answer'):
                    ws_questions.cell(row=row_num, column=col_offset, value=question['correct_answer'])
                    col_offset += 1
                
                # Add explanation
                if options.get('include_explanations'):
                    ws_questions.cell(row=row_num, column=col_offset, value=question.get('explanation', ''))
            
            # Auto-adjust column widths
            for ws in [ws_info, ws_questions]:
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column_letter].width = adjusted_width
            
            # Save to buffer
            buffer = io.BytesIO()
            wb.save(buffer)
            
            return {
                'content': buffer.getvalue(),
                'content_type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                'filename': f"{quiz.slug}_quiz.xlsx"
            }
            
        except ImportError:
            logger.error("openpyxl not installed for Excel export")
            raise ValueError("Excel export requires openpyxl package")
        except Exception as e:
            logger.error(f"Excel export failed: {str(e)}")
            raise
    
    def _export_to_html(self, quiz_data: Dict, quiz, **options) -> Dict[str, Any]:
        """Export quiz to HTML format"""
        # Render HTML template
        html_content = render_to_string('quizzes/export/quiz_export.html', {
            'quiz_data': quiz_data,
            'options': options,
            'export_date': timezone.now(),
        })
        
        return {
            'content': html_content.encode('utf-8'),
            'content_type': 'text/html',
            'filename': f"{quiz.slug}_quiz.html"
        }
    
    def _track_export(self, quiz, user, format_type: str, **options):
        """Track quiz export for analytics"""
        try:
            from .enhanced_models import QuizExport
            
            QuizExport.objects.create(
                quiz=quiz,
                user=user,
                format=format_type,
                include_answers=options.get('include_answers', True),
                include_explanations=options.get('include_explanations', True),
                include_analytics=options.get('include_analytics', False),
            )
        except Exception as e:
            logger.error(f"Failed to track export: {str(e)}")

# Global instance
quiz_export_service = QuizExportService()
